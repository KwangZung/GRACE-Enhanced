import os
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("❌ Vui lòng cài đặt thư viện bằng lệnh: pip install python-docx")
    exit()

doc = Document()

# Thêm Tiêu đề
heading = doc.add_heading('BÁO CÁO TIẾN ĐỘ NGHIÊN CỨU', 1)
doc.add_heading('CẢI TIẾN MÔ HÌNH GRACE TRONG PHÁT HIỆN LỖ HỔNG BẢO MẬT', 2)

# Thêm thông tin sinh viên
doc.add_paragraph('Kính gửi Thầy/Cô hướng dẫn,')
doc.add_paragraph('Sinh viên thực hiện: Nguyễn Hữu Hiếu')
doc.add_paragraph('Dưới đây là tóm tắt các điểm cải tiến cốt lõi về mặt kiến trúc hệ thống so với mô hình GRACE gốc, cùng với các khó khăn mà nhóm nghiên cứu đang tiến hành giải quyết trong giai đoạn chạy thực nghiệm.\n')

# Phần I
doc.add_heading('I. NHỮNG CẢI TIẾN CỐT LÕI VỀ KIẾN TRÚC', level=2)
doc.add_paragraph('Hệ thống mới (GRACE-Enhanced) đã khắc phục được các nhược điểm về tốc độ và độ chính xác của bản gốc thông qua 3 thay đổi lớn:')

doc.add_heading('1. Tối ưu hóa Đo lường Tương đồng Ngữ nghĩa (Semantic Similarity)', level=3)
doc.add_paragraph('Hạn chế của bản gốc: Sử dụng khoảng cách L2 Distance truyền thống và kỹ thuật T-SNE để nén chiều vector của CodeT5 từ 768 chiều xuống còn 2 chiều. Điều này gây mất mát lượng lớn thông tin không gian và khiến các vector phân bố không đồng đều.', style='List Bullet')
p = doc.add_paragraph('Giải pháp cải tiến:\n', style='List Bullet')
p.add_run('- Áp dụng kỹ thuật BERT Whitening: Giúp chuẩn hóa phân phối vector của CodeT5 (đẳng hướng hóa không gian) và giảm số chiều xuống phân nửa một cách khoa học, vừa tăng tốc độ xử lý vừa bảo toàn trọn vẹn đặc trưng ngữ nghĩa.\n')
p.add_run('- Thay thế L2 Distance bằng công cụ FAISS (Facebook AI Similarity Search): Một thư viện chuyên dụng giúp tính toán và truy xuất Top-K đoạn mã có độ tương đồng ngữ nghĩa cao nhất với tốc độ cực nhanh trên tập dữ liệu lớn.')

doc.add_heading('2. Nâng cấp Đo lường Tương đồng Cú pháp & Từ vựng (Syntactic & Lexical)', level=3)
doc.add_paragraph('Hạn chế của bản gốc: Sử dụng các công thức so sánh thủ công, đơn giản, chưa mang lại hiệu quả phân loại cao.', style='List Bullet')
p2 = doc.add_paragraph('Giải pháp cải tiến: Ứng dụng BM25 – một thuật toán xếp hạng siêu việt và tiêu chuẩn trong lĩnh vực Truy xuất Thông tin (Information Retrieval).\n', style='List Bullet')
p2.add_run('- Về Từ vựng (Lexical): Dùng công cụ phân tích tĩnh Joern để làm sạch đoạn mã (loại bỏ ký tự thừa, chuẩn hóa tên biến). Sau đó, áp dụng BM25 lên đầu ra này để đo lường độ tương đồng mặt chữ.\n')
p2.add_run('- Về Cú pháp (Syntactic): Dùng Joern trích xuất biểu diễn Cây cú pháp trừu tượng (AST). Để giảm tải độ phức tạp khi tính toán trực tiếp trên đồ thị cây, hệ thống áp dụng kỹ thuật SIM SBT (Structure-Based Traversal) để duỗi cây AST thành một chuỗi tuyến tính. Cuối cùng, áp dụng BM25 theo từng cụm N-gram trên chuỗi này để đo lường cấu trúc cú pháp.')

doc.add_heading('3. Cơ chế Khái quát hóa với RAG Few-shot Prompting', level=3)
doc.add_paragraph('Sau khi tổng hòa 3 độ đo (Semantic, Lexical, Syntactic) bằng một trọng số để tìm ra đoạn code mẫu (Example) giống với code mục tiêu (Target) nhất, hệ thống không kết luận ngay.', style='List Bullet')
doc.add_paragraph('Áp dụng LLM (Large Language Model): Đoạn code mẫu cùng với nhãn thực tế của nó sẽ được đưa vào làm mồi (In-context Learning) cho LLM đánh giá đoạn code mục tiêu.', style='List Bullet')
doc.add_paragraph('Kỹ thuật RAG Few-shot (Retrieval-Augmented Generation): Thay vì để LLM phán đoán mù mờ (Zero-shot), RAG cung cấp một "hệ quy chiếu" động. Việc này giúp LLM học được cách đối chiếu luồng dữ liệu (Data Flow) giữa code có lỗi và code an toàn để đưa ra quyết định cuối cùng chính xác nhất.', style='List Bullet')

# Phần II
doc.add_heading('II. NHỮNG KHÓ KHĂN & THÁCH THỨC HIỆN TẠI', level=2)
doc.add_paragraph('Trong quá trình triển khai cấu hình và chạy thực nghiệm, hệ thống đang phải đối mặt và xử lý các vấn đề sau:')
doc.add_paragraph('1. Giới hạn Tài nguyên API (Rate Limits/Quotas): Các mô hình LLM hiện đại có năng lực suy luận cao (như Gemini/GPT) thường giới hạn rất khắt khe số lượng yêu cầu (Request) đối với các tài khoản thử nghiệm miễn phí (ví dụ: cạn kiệt Quota chỉ sau 20 mẫu test/ngày), làm gián đoạn quá trình đánh giá quy mô lớn.')
doc.add_paragraph('2. Đánh đổi giữa Tốc độ và Sự tập trung của LLM (Attention Dilution): Để lách luật giới hạn API, nhóm đã thử nghiệm kỹ thuật Batch Prompting (Gộp nhiều đoạn mã mục tiêu vào cùng một lần hỏi LLM). Tuy nhiên, việc nhồi nhét quá nhiều cấu trúc đồ thị (Node/Edge) vào một cửa sổ ngữ cảnh khiến cơ chế Attention của LLM bị loãng, dẫn đến việc phán đoán luồng dữ liệu đôi lúc bị nhầm lẫn giữa các hàm với nhau.')
doc.add_paragraph('3. Hướng giải quyết: Nhóm đang tiến hành thiết lập cơ sở hạ tầng đám mây (Google Cloud Project) có liên kết thanh toán để mở khóa hạn mức API, đưa Batch Size về mức tối ưu (so sánh từng hàm một), kết hợp các kỹ thuật phân tích Regex chuyên sâu để ép LLM trả về kết quả tuân thủ nghiêm ngặt định dạng đầu ra, đảm bảo tính toàn vẹn của dữ liệu thực nghiệm.')

# Lưu file
doc.save('Bao_Cao_Tien_Do_GRACE.docx')
print("✅ Đã tạo thành công file Bao_Cao_Tien_Do_GRACE.docx. Ông có thể mở lên xem và gửi đi!")